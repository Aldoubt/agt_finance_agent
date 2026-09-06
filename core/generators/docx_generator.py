"""Generate the compact product-photo and proof document."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}


def _set_cell_margins(cell, *, top=45, start=55, bottom=45, end=55) -> None:
    """Use compact cell padding (twips) so three proof cards fit comfortably."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _fit_image(path: Path, *, max_width: float, max_height: float) -> tuple[float, float]:
    """Return an undistorted image size in inches bounded by width and height."""
    try:
        with Image.open(path) as image:
            oriented = ImageOps.exif_transpose(image)
            width_px, height_px = oriented.size
    except Exception:
        return max_width, max_height

    if width_px <= 0 or height_px <= 0:
        return max_width, max_height
    ratio = width_px / height_px
    width = min(max_width, max_height * ratio)
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    return width, height


def _add_picture(paragraph, path: Path, *, max_width: float, max_height: float) -> bool:
    if not path.exists() or path.suffix.lower() not in IMAGE_SUFFIXES:
        return False
    width, height = _fit_image(path, max_width=max_width, max_height=max_height)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    return True


class DocxGenerator:
    PHOTO_MAX_WIDTH = 5.8
    PHOTO_MAX_HEIGHT = 5.65
    PROOF_CARD_MAX_WIDTH = 2.05
    PROOF_CARD_IMAGE_BUDGET = 3.15
    PROOFS_PER_PAGE = 6

    def _configure_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

        normal = doc.styles["Normal"]
        normal.font.size = Pt(9)
        normal.paragraph_format.space_after = Pt(3)

        heading1 = doc.styles["Heading 1"]
        heading1.font.size = Pt(15)
        heading1.paragraph_format.space_before = Pt(4)
        heading1.paragraph_format.space_after = Pt(4)
        # Prevent a Heading 1 -> Heading 2 -> tall image keep-chain from moving
        # the whole first photo section to page 2 and leaving page 1 mostly empty.
        heading1.paragraph_format.keep_with_next = False

        heading2 = doc.styles["Heading 2"]
        heading2.font.size = Pt(11)
        heading2.paragraph_format.space_before = Pt(4)
        heading2.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _photo_groups(items) -> tuple[dict[str, list[str]], list[str], list[str]]:
        groups: dict[str, list[str]] = {}
        order: list[str] = []
        missing_items: list[str] = []
        for item in items:
            if not item.photo_files:
                missing_items.append(item.name)
                continue
            for photo in item.photo_files:
                if photo not in groups:
                    groups[photo] = []
                    order.append(photo)
                if item.name not in groups[photo]:
                    groups[photo].append(item.name)
        return groups, order, missing_items

    @staticmethod
    def _proof_groups(items) -> list[dict]:
        groups: dict[str, list] = {}
        order: list[str] = []
        for item in items:
            invoice_key = item.invoice_files[0] if item.invoice_files else f"item:{item.name}"
            if invoice_key not in groups:
                groups[invoice_key] = []
                order.append(invoice_key)
            groups[invoice_key].append(item)

        cards: list[dict] = []
        for invoice_key in order:
            grouped_items = groups[invoice_key]
            proofs: list[str] = []
            payments: list[str] = []
            for item in grouped_items:
                for proof in item.purchase_proof_files:
                    if proof not in proofs:
                        proofs.append(proof)
                for payment in item.payment_files:
                    if payment not in payments:
                        payments.append(payment)
            cards.append(
                {
                    "names": [item.name for item in grouped_items],
                    "proofs": proofs,
                    "payments": payments,
                }
            )
        return cards

    def _render_proof_card(self, cell, card: dict, index: int) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(cell)

        title = cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        run = title.add_run(f"{index:02d} " + " + ".join(card["names"]))
        run.bold = True
        run.font.size = Pt(8.5)

        image_entries: list[tuple[str, str]] = []
        image_entries.extend(("购买凭证", path) for path in card["proofs"])
        image_entries.extend(("支付凭证", path) for path in card["payments"])
        image_count = max(1, len(image_entries))
        each_height = min(self.PROOF_CARD_IMAGE_BUDGET, self.PROOF_CARD_IMAGE_BUDGET / image_count)

        if not image_entries:
            p = cell.add_paragraph("（无购买凭证）")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(8)
            return

        previous_label = None
        for label, value in image_entries:
            if label != previous_label or label == "支付凭证":
                label_p = cell.add_paragraph()
                label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label_p.paragraph_format.space_before = Pt(0)
                label_p.paragraph_format.space_after = Pt(1)
                label_run = label_p.add_run(label)
                label_run.bold = True
                label_run.font.size = Pt(7.5)
            path = Path(value)
            picture_p = cell.add_paragraph()
            if not _add_picture(
                picture_p,
                path,
                max_width=self.PROOF_CARD_MAX_WIDTH,
                max_height=each_height,
            ):
                picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                text_run = picture_p.add_run(path.name)
                text_run.font.size = Pt(7)
            previous_label = label

    def _render_proof_pages(self, doc: Document, cards: list[dict]) -> None:
        if not cards:
            doc.add_paragraph("（无购买凭证）")
            return

        for page_start in range(0, len(cards), self.PROOFS_PER_PAGE):
            if page_start:
                doc.add_page_break()
            page_cards = cards[page_start : page_start + self.PROOFS_PER_PAGE]
            table = doc.add_table(rows=2, cols=3)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for column in table.columns:
                column.width = Inches(2.35)
            for row in table.rows:
                _prevent_row_split(row)
                for cell in row.cells:
                    cell.width = Inches(2.35)
                    _set_cell_margins(cell)

            for offset, card in enumerate(page_cards):
                row_index, column_index = divmod(offset, 3)
                self._render_proof_card(
                    table.cell(row_index, column_index),
                    card,
                    page_start + offset + 1,
                )

    def generate(self, items, output_path: str | Path) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_document(doc)
        doc.add_heading("采购归档", level=1)
        count_p = doc.add_paragraph(f"商品数量：{len(items)}")
        count_p.paragraph_format.space_after = Pt(2)
        doc.add_heading("一、商品名称与实物图", level=1)

        photo_groups, photo_order, missing_items = self._photo_groups(items)
        for index, photo in enumerate(photo_order, 1):
            names = " + ".join(photo_groups[photo])
            heading = doc.add_heading(f"{index:02d} {names}", level=2)
            heading.paragraph_format.keep_with_next = True
            path = Path(photo)
            picture_p = doc.add_paragraph()
            if not _add_picture(
                picture_p,
                path,
                max_width=self.PHOTO_MAX_WIDTH,
                max_height=self.PHOTO_MAX_HEIGHT,
            ):
                picture_p.add_run(f"（图片文件不存在：{path.name}）")

        if missing_items:
            doc.add_heading("未关联实物图", level=2)
            for name in missing_items:
                doc.add_paragraph(name)

        # Start the compact proof gallery on a clean page. The removed blank
        # page was before section 1, not this intentional section separator.
        doc.add_page_break()
        doc.add_heading("二、购买凭证与支付凭证", level=1)
        self._render_proof_pages(doc, self._proof_groups(items))

        doc.save(output)
        return output

