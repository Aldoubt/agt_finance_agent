from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from core.generators.docx_generator import DocxGenerator
from core.graph.purchase_graph import PurchaseItem


def _image(path: Path, size=(600, 1200)) -> str:
    Image.new("RGB", size, "white").save(path)
    return str(path)


def _items(tmp_path: Path, count: int) -> list[PurchaseItem]:
    rows = []
    for index in range(count):
        photo = _image(tmp_path / f"photo_{index}.png")
        proof = _image(tmp_path / f"proof_{index}.png", size=(400, 900))
        invoice = tmp_path / f"invoice_{index}.pdf"
        invoice.write_bytes(b"%PDF-demo")
        rows.append(
            PurchaseItem(
                name=f"商品{index + 1}",
                invoice_files=[str(invoice)],
                photo_files=[photo],
                purchase_proof_files=[proof],
            )
        )
    return rows


def test_docx_uses_six_proof_cards_per_page(tmp_path):
    output = DocxGenerator().generate(_items(tmp_path, 7), tmp_path / "archive.docx")
    doc = Document(output)

    assert len(doc.tables) == 2
    assert all(len(table.rows) == 2 for table in doc.tables)
    assert all(len(row.cells) == 3 for table in doc.tables for row in table.rows)
    assert "01 商品1" in doc.tables[0].cell(0, 0).text
    assert "06 商品6" in doc.tables[0].cell(1, 2).text
    assert "07 商品7" in doc.tables[1].cell(0, 0).text


def test_docx_limits_tall_photo_height_and_has_no_cover_page_break(tmp_path):
    output = DocxGenerator().generate(_items(tmp_path, 1), tmp_path / "archive.docx")
    doc = Document(output)

    assert doc.inline_shapes[0].height <= Inches(DocxGenerator.PHOTO_MAX_HEIGHT)
    # With one proof page there is only the intentional break before section 2;
    # there is no artificial cover-page break before the photo section.
    page_breaks = doc._element.xml.count('w:type="page"')
    assert page_breaks == 1


def test_docx_keeps_payment_in_same_proof_card(tmp_path):
    items = _items(tmp_path, 1)
    payment = _image(tmp_path / "payment.png", size=(400, 900))
    items[0].payment_files = [payment]
    output = DocxGenerator().generate(items, tmp_path / "archive.docx")
    doc = Document(output)

    card_text = doc.tables[0].cell(0, 0).text
    assert "购买凭证" in card_text
    assert "支付凭证" in card_text
    assert len(doc.tables) == 1
